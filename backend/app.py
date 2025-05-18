from fastapi import FastAPI, HTTPException, Form, File, UploadFile, Depends, BackgroundTasks , Request, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse , Response
from pydantic import BaseModel
from typing import List, Dict, Optional, Any
from uuid import uuid4
import ollama
import requests
import json
import os
import PyPDF2
import docx
import pandas as pd
import io
from io import BytesIO, StringIO
import numpy as np
import shutil
from sentence_transformers import SentenceTransformer
import faiss
import pickle
from deep_translator import GoogleTranslator
import datetime
import re
import time
import requests
import sympy as sp
from sympy.parsing.sympy_parser import parse_expr
import json
import matplotlib.pyplot as plt
import base64
from diffusers import StableDiffusionPipeline , DPMSolverMultistepScheduler
import torch

# Chat history file for persistence
CHAT_HISTORY_FILE = "chats.json"
VECTOR_DB_DIR = "vector_db"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
TOP_K_RESULTS = 5



# Create vector db directory if it doesn't exist
os.makedirs(VECTOR_DB_DIR, exist_ok=True)

# Load Stable Diffusion model once at startup (do this outside the route)
# Load the pipeline with advanced sampler for sharper outputs
image_pipe = StableDiffusionPipeline.from_pretrained(
    "runwayml/stable-diffusion-v1-5", 
    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
)

# Set the sampler to DPM++ 2M Karras for better quality
image_pipe.scheduler = DPMSolverMultistepScheduler.from_config(image_pipe.scheduler.config)

image_pipe = image_pipe.to("cuda" if torch.cuda.is_available() else "cpu")


# Load chats from file on startup
def load_chats_from_file():
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, "r") as f:
            data = json.load(f)
            return [Chat(**chat) for chat in data]
    return []

# Save chats to file after any modification
def save_chats_to_file():
    try:
        # Convert chats to dict for serialization
        chats_data = [chat.dict() for chat in chats]
        
        # Write to file
        with open("chats.json", "w") as f:
            json.dump(chats_data, f, indent=2)
    except Exception as e:
        print(f"Error saving chats: {e}")

# Initialize FastAPI app
app = FastAPI()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# SerpAPI keys
SERPAPI_KEYS = [
    "736c0c4229125891db11c1f0cdfe9485dfc96115dc53237931290a001a96c7ea",
    "29f2ca31b18eb58a4a3d9bbff7e750e4f4ade20c0db8cef364363aa83543ec61",
]

# Supported languages and their codes for translation
SUPPORTED_LANGUAGES = {
    "English": "en",
    "Hindi": "hi",
    "Tamil": "ta",  
    "Telugu": "te",   
    "Spanish": "es",
    "French": "fr",
    "German": "de",
    "Japanese": "ja",
    "Chinese": "zh-CN",
    "Russian": "ru",
    "Arabic": "ar",
    "Portuguese": "pt"
}

# Define the AI's name and introduction phrases for common languages
AI_NAME = "Xendrix"
AI_INTRODUCTIONS = {
    "English": f"My name is {AI_NAME}. I'm an AI assistant designed to help you with various tasks.",
    "Hindi": f"मेरा नाम {AI_NAME} है। मैं आपकी सहायता के लिए डिज़ाइन किया गया एक AI सहायक हूँ।",
    "Tamil": f"என் பெயர் {AI_NAME}. நான் பல்வேறு பணிகளில் உங்களுக்கு உதவ வடிவமைக்கப்பட்ட ஒரு AI உதவியாளர்.",
    "Telugu": f"నా పేరు {AI_NAME}. నేను వివిధ పనులలో మీకు సహాయపడటానికి రూపొందించబడిన AI సహాయకుడిని.",
    "Spanish": f"Mi nombre es {AI_NAME}. Soy un asistente de IA diseñado para ayudarte con varias tareas.",
    "French": f"Je m'appelle {AI_NAME}. Je suis un assistant IA conçu pour vous aider dans diverses tâches.",
    "German": f"Mein Name ist {AI_NAME}. Ich bin ein KI-Assistent, der Ihnen bei verschiedenen Aufgaben hilft.",
    "Japanese": f"私の名前は{AI_NAME}です。私はさまざまなタスクをお手伝いするためのAIアシスタントです。",
    "Chinese": f"我的名字是{AI_NAME}。我是一个设计用来帮助你完成各种任务的AI助手。",
    "Russian": f"Меня зовут {AI_NAME}. Я ИИ-ассистент, созданный, чтобы помогать вам с различными задачами.",
    "Arabic": f"اسمي {AI_NAME}. أنا مساعد ذكاء اصطناعي مصمم لمساعدتك في المهام المختلفة.",
    "Portuguese": f"Meu nome é {AI_NAME}. Sou um assistente de IA projetado para ajudá-lo com várias tarefas."
}

# Name detection patterns for various languages
NAME_PATTERNS = {
    "English": [
        r"what(?:'s| is) your name",
        r"who are you",
        r"introduce yourself",
        r"tell me your name",
        r"what should I call you"
    ],
    "Hindi": [
        r"(?:तुम्हारा|आपका|तेरा) नाम क्या है",
        r"(?:तुम|आप) कौन हो",
        r"(?:अपना पर?चय|परिचय) दो", 
        r"(?:अपना नाम बताओ|बताईए)",
        r"मैं (?:तुम्हें|आपको) क्या कहूँ"
    ],
    "Tamil": [
        r"உன்னுடைய பெயர் என்ன",
        r"உங்கள் பெயர் என்ன",
        r"நீ யார்",
        r"நீங்கள் யார்",
        r"உன்னை அறிமுகப்படுத்திக்கொள்",
        r"உன் பெயரை சொல்"
    ],
    "Telugu": [
        r"నీ పేరు ఏమిటి",
        r"మీ పేరు ఏమిటి",
        r"నువ్వు ఎవరు",
        r"మీరు ఎవరు",
        r"నిన్ను నువ్వు పరిచయం చేసుకో",
        r"నీ పేరు చెప్పు"
    ],
    "Spanish": [
        r"(?:cuál es tu nombre|cómo te llamas)",
        r"quién eres",
        r"preséntate",
        r"dime tu nombre",
        r"cómo debo llamarte"
    ],
    "French": [
        r"(?:quel est ton nom|comment t'appelles-tu|comment tu t'appelles)",
        r"qui es-tu",
        r"présente-toi",
        r"dis-moi ton nom",
        r"comment dois-je t'appeler"
    ],
    "German": [
        r"(?:wie ist dein Name|wie heißt du)",
        r"wer bist du",
        r"stelle dich vor",
        r"sag mir deinen Namen",
        r"wie soll ich dich nennen"
    ]
}

# Function to detect if a message is asking for the AI's name
def is_asking_name(message, language="English"):
    # Default to English patterns if language not in our patterns
    patterns = NAME_PATTERNS.get(language, NAME_PATTERNS["English"])
    
    # Convert message to lowercase for case-insensitive matching
    message_lower = message.lower()
    
    # Check if any pattern matches
    for pattern in patterns:
        if re.search(pattern, message_lower):
            return True
    
    # If no match in the specific language, try English patterns as fallback
    if language != "English":
        for pattern in NAME_PATTERNS["English"]:
            if re.search(pattern, message_lower):
                return True
    
    return False

# Data models
class Message(BaseModel):
    sender: str
    text: str
    language: str = "English"  # Default language
    image_path: Optional[str] = None 

class Chat(BaseModel):
    id: str
    name: str = "Untitled Chat"
    messages: List[Message]
    document_id: Optional[str] = None
    file_content: Optional[str] = None
    file_name: Optional[str] = None
    language: str = "English"  # Default language for the chat

class ChatRequest(BaseModel):
    message: str
    language: str = "English"
    
class Document(BaseModel):
    id: str
    name: str
    chunks: List[str]
    
class DocumentInfo(BaseModel):
    id: str
    name: str
    
class DocumentMetadata(BaseModel):
    documents: List[DocumentInfo] = []

class VisualizationRequest(BaseModel):
    query: str
    chat_id: Optional[str] = None
    document_id: Optional[str] = None
    chart_type: str = "bar"  # Default chart type
    language: str = "English"

class DataPoint(BaseModel):
    label: str
    value: float

class VisualizationData(BaseModel):
    title: str
    data: List[DataPoint]
    chart_type: str
    x_label: str = "Categories"
    y_label: str = "Values"

# Load initial chat state
chats: List[Chat] = load_chats_from_file()

# Initialize sentence transformer model for embeddings
model = SentenceTransformer('all-MiniLM-L6-v2')

# Load document metadata
def load_document_metadata():
    metadata_path = os.path.join(VECTOR_DB_DIR, "metadata.json")
    if os.path.exists(metadata_path):
        with open(metadata_path, "r") as f:
            data = json.load(f)
            return DocumentMetadata(**data)
    return DocumentMetadata()

# Save document metadata
def save_document_metadata(metadata: DocumentMetadata):
    metadata_path = os.path.join(VECTOR_DB_DIR, "metadata.json")
    with open(metadata_path, "w") as f:
        json.dump(metadata.dict(), f, indent=2)

# Load document metadata on startup
document_metadata = load_document_metadata()

# Translation functions
# First, let's modify the translation function to be more robust
def translate_text(text, source_lang, target_lang, max_retries=3):
    """Translate text with retry mechanism and fallback options"""
    # Skip translation if languages are the same
    if source_lang == target_lang or source_lang == "English" and target_lang == "en" or target_lang == "English" and source_lang == "en":
        return text
    
    # Normalize language codes
    if source_lang in SUPPORTED_LANGUAGES:
        source_lang = SUPPORTED_LANGUAGES[source_lang]
    if target_lang in SUPPORTED_LANGUAGES:
        target_lang = SUPPORTED_LANGUAGES[target_lang]
    
    # Multiple translation services to try
    translation_methods = [
        # Method 1: Google Translator with longer timeout
        lambda: GoogleTranslator(source=source_lang if source_lang != "auto" else "auto", 
                               target=target_lang, timeout=10).translate(text),
        
        # Method 2: Alternative approach with direct requests
        lambda: _fallback_translate(text, source_lang, target_lang)
    ]
    
    # Try each method with retries
    for method in translation_methods:
        for attempt in range(max_retries):
            try:
                return method()
            except Exception as e:
                print(f"Translation error (attempt {attempt+1}/{max_retries}): {e}")
                if attempt == max_retries - 1 and method == translation_methods[-1]:
                    print("All translation methods failed. Returning original text.")
                    return text  # Return original text as last resort
                time.sleep(1)  # Short delay before retry
    
    return text  # Fallback to original text

def _fallback_translate(text, source_lang, target_lang):
    """Alternative translation method using requests"""
    try:
        # Use a different translation endpoint
        url = "https://translate.googleapis.com/translate_a/single"
        params = {
            "client": "gtx",
            "sl": source_lang,
            "tl": target_lang,
            "dt": "t",
            "q": text
        }
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        
        response = requests.get(url, params=params, headers=headers, timeout=15)
        if response.status_code == 200:
            # Parse the response
            result = response.json()
            translated_text = ""
            for sentence in result[0]:
                if sentence[0]:
                    translated_text += sentence[0]
            return translated_text
        return text
    except Exception as e:
        print(f"Fallback translation error: {e}")
        return text
    
def extract_data_for_visualization(text: str) -> Dict[str, Any]:
    """
    Extract data from text that can be used for visualization.
    Always returns a dictionary with "labels" and "values" keys.
    """
    import re
    from typing import Dict, Any, List
    
    # Initialize default return structure
    result = {
        "labels": [],
        "values": []
    }
    
    try:
        # Check for table format with pipe separators
        if '|' in text:
            lines = [line.strip() for line in text.split('\n') if '|' in line]
            
            # Need at least a few lines for a table
            if len(lines) < 2:
                return result
                
            # Extract header to identify columns
            header_line = lines[0]
            header_parts = [part.strip().lower() for part in header_line.split('|') if part.strip()]
            
            # Skip separator lines like |---|---|---|
            data_lines = [line for line in lines if not re.match(r'\s*\|[\s\-:]+\|\s*$', line)]
            if len(data_lines) > 0:
                data_lines = data_lines[1:]  # Skip header row
            
            # If no data lines left, return empty result
            if not data_lines:
                return result
                
            # We need at least 2 columns for label-value pairs
            if len(header_parts) < 2:
                return result
                
            # Assume first column is the label column and the second is the value column
            label_col_index = 0
            value_col_index = 1
            
            # But look for more appropriate columns based on headers
            for i, header in enumerate(header_parts):
                header_lower = header.lower()
                # Check if this looks like a label column (year, date, category, etc.)
                if any(term in header_lower for term in ['year', 'date', 'month', 'category', 'model', 'name', 'product']):
                    label_col_index = i
                    # Look for value column after finding label column
                    for j, value_header in enumerate(header_parts):
                        value_header_lower = value_header.lower()
                        if j != i and any(term in value_header_lower for term in 
                                         ['value', 'sales', 'amount', 'number', 'quantity', 'total', 'sum', 'count']):
                            value_col_index = j
                            break
                    break
            
            # Extract data points
            data_points = []
            for line in data_lines:
                parts = [part.strip() for part in line.split('|') if part.strip()]
                
                # Make sure we have enough parts
                if len(parts) <= max(label_col_index, value_col_index):
                    continue
                
                # Get label
                label = parts[label_col_index]
                
                # Get value and clean it
                value_text = parts[value_col_index]
                value_match = re.search(r'(\d+(?:,\d+)*(?:\.\d+)?)', value_text)
                if value_match:
                    # Remove commas from number string
                    clean_value = value_match.group(1).replace(',', '')
                    try:
                        value = float(clean_value)
                        # Only add if the label isn't a number that equals the value
                        # This prevents adding entries where year is both label and value
                        try:
                            label_as_num = float(label)
                            if abs(label_as_num - value) > 0.01:  # Use a small epsilon for float comparison
                                data_points.append((label, value))
                        except ValueError:
                            # Label isn't a number, so add the point
                            data_points.append((label, value))
                    except ValueError:
                        pass
            
            if data_points:
                result = {
                    "labels": [point[0] for point in data_points],
                    "values": [point[1] for point in data_points]
                }
                return result
            
        # Pattern 1: Look for tables with numbers
        table_pattern = re.compile(r'([\w\s]+)[\s\|:]+(\d+(?:,\d+)*(?:\.\d+)?)')
        matches = table_pattern.findall(text)
        if matches:
            data_points = []
            for label, value_str in matches:
                # Remove commas from number strings
                clean_value = value_str.replace(',', '')
                try:
                    value = float(clean_value)
                    # Avoid adding entries where label equals value (like year-as-year)
                    try:
                        label_as_num = float(label.strip())
                        if abs(label_as_num - value) > 0.01:
                            data_points.append((label.strip(), value))
                    except ValueError:
                        data_points.append((label.strip(), value))
                except ValueError:
                    pass
                    
            if data_points:
                result = {
                    "labels": [point[0] for point in data_points],
                    "values": [point[1] for point in data_points]
                }
                return result
        
        # Pattern 2: Look for lists with numbers
        list_pattern = re.compile(r'[-*•]\s*([\w\s]+):\s*(\d+(?:,\d+)*(?:\.\d+)?)')
        matches = list_pattern.findall(text)
        if matches:
            data_points = []
            for label, value_str in matches:
                # Remove commas from number strings
                clean_value = value_str.replace(',', '')
                try:
                    value = float(clean_value)
                    # Avoid adding entries where label equals value
                    try:
                        label_as_num = float(label.strip())
                        if abs(label_as_num - value) > 0.01:
                            data_points.append((label.strip(), value))
                    except ValueError:
                        data_points.append((label.strip(), value))
                except ValueError:
                    pass
                    
            if data_points:
                result = {
                    "labels": [point[0] for point in data_points],
                    "values": [point[1] for point in data_points]
                }
                return result
        
        # Pattern 3: Look for sentences with percentages or numbers
        sentence_pattern = re.compile(r'([\w\s]+)(?:is|was|at|:)\s*(\d+(?:,\d+)*(?:\.\d+)?)\s*(?:%|percent)?')
        matches = sentence_pattern.findall(text)
        if matches:
            data_points = []
            for label, value_str in matches:
                # Remove commas from number strings
                clean_value = value_str.replace(',', '')
                try:
                    value = float(clean_value)
                    # Avoid adding entries where label equals value
                    try:
                        label_as_num = float(label.strip())
                        if abs(label_as_num - value) > 0.01:
                            data_points.append((label.strip(), value))
                    except ValueError:
                        data_points.append((label.strip(), value))
                except ValueError:
                    pass
                    
            if data_points:
                result = {
                    "labels": [point[0] for point in data_points],
                    "values": [point[1] for point in data_points]
                }
                return result
        
        # Return the default empty result if no patterns matched
        return result
    
    except Exception as e:
        print(f"Error extracting data for visualization: {e}")
        # Return empty but valid structure on error
        return result
    
def generate_visualization(data: Dict[str, Any], chart_type: str = "bar") -> Optional[str]:
    """
    Generate a visualization based on the provided data and chart type.
    Returns a base64 encoded image.
    """
    try:
        labels = data["labels"]
        values = data["values"]
        
        # Handle pie chart completely separately to avoid any axis issues
        if chart_type == "pie":
            # Create a new figure with appropriate size for pie
            plt.figure(figsize=(10, 8))
            
            # Only use top 7 items for pie charts, group the rest as "Other"
            if len(labels) > 7:
                top_labels = labels[:6]
                top_values = values[:6]
                other_value = sum(values[6:])
                pie_labels = top_labels + ["Other"]
                pie_values = top_values + [other_value]
            else:
                pie_labels = labels
                pie_values = values
                
            # Create pie chart with nice colors
            plt.pie(pie_values, labels=None, autopct='%1.1f%%', startangle=90, 
                   shadow=False, colors=plt.cm.tab10.colors)
            
            # Add legend outside the pie for better readability
            plt.legend(pie_labels, loc="center left", bbox_to_anchor=(1, 0.5))
            
            # Force equal aspect ratio for circular pie
            plt.axis('equal')
            
            # Add title for pie chart
            plt.title("Pie Chart", fontsize=14)
            
            # Tight layout for pie chart
            plt.tight_layout()
            
        else:
            # Create a new figure with appropriate size for other charts
            plt.figure(figsize=(12, 7))
            
            # Handle long labels better
            max_label_length = max(len(str(label)) for label in labels)
            
            # Calculate positions for x-ticks
            positions = range(len(labels))
            
            # Generate chart based on type
            if chart_type == "bar":
                # Use a visually appealing color
                bars = plt.bar(positions, values, color='#1f77b4')
                
                # Add data labels on top of bars for important values
                if len(labels) <= 10:  # Only add data labels if not too many bars
                    for bar in bars:
                        height = bar.get_height()
                        plt.text(bar.get_x() + bar.get_width()/2., height + 0.01*max(values),
                                f'{height:,.0f}', 
                                ha='center', va='bottom', fontsize=9)
            
            elif chart_type == "line":
                plt.plot(positions, values, marker='o', linestyle='-', color='#1f77b4', linewidth=2)
                
                # Add data points labels if not too many
                if len(labels) <= 10:
                    for i, val in enumerate(values):
                        plt.text(i, val + 0.01*max(values), f'{val:,.0f}', 
                                ha='center', va='bottom', fontsize=9)
            
            elif chart_type == "scatter":
                plt.scatter(positions, values, s=100, alpha=0.7, color='#1f77b4')
                
                # Add connecting line for trend visibility
                plt.plot(positions, values, '--', color='#1f77b4', alpha=0.5)
                
                # Add data labels if not too many points
                if len(labels) <= 10:
                    for i, val in enumerate(values):
                        plt.text(i, val + 0.01*max(values), f'{val:,.0f}', 
                                ha='center', va='bottom', fontsize=9)
            else:
                # Default to bar chart
                plt.bar(positions, values, color='#1f77b4')
            
            # Add labels and title for non-pie charts
            plt.xlabel("Categories", fontsize=12)
            plt.ylabel("Values", fontsize=12)
            plt.title(f"{chart_type.capitalize()} Chart", fontsize=14)
            
            # Set x-tick positions and labels
            plt.xticks(positions, labels)
            
            # Adjust x-tick labels rotation based on label length
            if max_label_length > 8:
                rotation = 45
                ha = 'right'
            else:
                rotation = 0
                ha = 'center'
                
            plt.xticks(rotation=rotation, ha=ha)
            
            # Add grid for better readability in most chart types
            plt.grid(axis='y', linestyle='--', alpha=0.7)
            
            # Adjust layout with more padding for rotated labels
            plt.tight_layout(pad=3.0 if rotation > 0 else 1.5)
        
        # Save the plot to a bytes buffer
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)
        
        # Convert to base64 for embedding in HTML/response
        img_str = base64.b64encode(buf.read()).decode('utf-8')
        plt.close()
        
        return img_str
        
    except Exception as e:
        print(f"Error generating visualization: {e}")
        return None

def is_math_problem(message: str) -> bool:
    """
    Simple function to detect if a message is a math problem.
    Returns True if it's a math problem, False otherwise.
    """
    message_lower = message.lower()
    
    # Math patterns to check
    math_patterns = [
        # Explicit equations
        r'\b[xyz]\s*=\s*[\d\+\-\*/\(\)\.\^]+',
        r'[\d\+\-\*/\(\)\.]+\s*=\s*[\d\+\-\*/\(\)\.x-z]+',
        # Math operators with numbers
        r'\d+\s*[\+\-\*/\^]\s*\d+',
    ]
    
    # Math keywords that strongly indicate a math problem
    strong_math_keywords = [
        "solve", "calculate", "equation", "formula", "algebra", 
        "find x", "find y", "value of", "equal to"
    ]
    
    # Regular math keywords
    math_keywords = [
        "math", "calculation", "price", "cost", "percent", "ratio", 
        "average", "mean", "median", "probability", "discount"
    ]
    
    # Check for definite math patterns first
    for pattern in math_patterns:
        if re.search(pattern, message):
            return True
            
    # Check for strong math keywords
    if any(keyword in message_lower for keyword in strong_math_keywords):
        return True
        
    # Check for combination of regular math keywords and length
    math_keyword_count = sum(1 for keyword in math_keywords if keyword in message_lower)
    if math_keyword_count >= 2 and len(message.split()) > 5:
        return True
        
    return False

def is_weather_query(message: str) -> bool:
    """Simple function to detect weather queries"""
    message_lower = message.lower()
    weather_keywords = ["weather", "temperature", "forecast", "rain", "sunny", "hot", "cold", "humidity"]
    
    return any(keyword in message_lower for keyword in weather_keywords)

def extract_location(message: str) -> Optional[str]:
    """Extract location from a weather query"""
    location_patterns = [
        r"(?:weather|temperature|forecast)(?:\s+in|\s+at|\s+for|\s+of)\s+([\w\s]+?)(?:$|\?|\.)",
        r"(?:how(?:'s| is) the|what(?:'s| is) the)\s+(?:weather|temperature)(?:\s+in|\s+at)?\s+([\w\s]+?)(?:$|\?|\.)",
        r"(?:in|at|for)\s+([\w\s]+?)(?:'s|\s+weather|\s+temperature)(?:$|\?|\.)"
    ]
    
    for pattern in location_patterns:
        match = re.search(pattern, message, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return None

def format_weather_response(weather_info: dict, original_query: str, language: str) -> str:
    """Format weather information into a natural response"""
    if "error" in weather_info:
        response = f"I'm sorry, I couldn't get weather information. {weather_info['error']}"
    else:
        response = f"The current weather in {weather_info['location']} is {weather_info['temperature']} with {weather_info['condition']}. Humidity is {weather_info['humidity']} and wind speed is {weather_info['wind_speed']}."
    
    # Translate if needed
    if language != "English" and language in SUPPORTED_LANGUAGES:
        response = translate_text(response, "en", language)
        
    return response

async def get_regular_chat_response(message: str, context: str, language: str) -> str:
    """Get response for regular chat using Mistral"""
    try:
        print("Using Mistral for regular conversation")
        full_prompt = f"{context}\n\nUser asked: {message}\nPlease respond based on the conversation context. Remember to introduce yourself as {AI_NAME} if appropriate."
        response = ollama.chat(model="mistral", messages=[{"role": "user", "content": full_prompt}])
        ai_answer = response['message']['content']

        # Check if response contains code
        if any(keyword in message.lower() for keyword in ["code", "function", "script", "program", "class"]):
            if "```" not in ai_answer:
                ai_answer = f"```python\n{ai_answer.strip()}\n```"
        
        # Translate if needed
        if language != "English" and language in SUPPORTED_LANGUAGES:
            if "```" in ai_answer:
                # Handle code blocks during translation
                parts = ai_answer.split("```")
                for i in range(0, len(parts), 2):  # Translate only non-code parts
                    parts[i] = translate_text(parts[i], "en", SUPPORTED_LANGUAGES[language])
                ai_answer = "```".join(parts)
            else:
                ai_answer = translate_text(ai_answer, "en", language)
                
        return ai_answer
    except Exception as e:
        error_msg = f"Error generating response: {str(e)}"
        print(error_msg)
        return "I'm sorry, I encountered an error while processing your request."

def get_weather_information(location: str, unit: str = "metric") -> dict:
    """
    Fetches current weather information for a given location using the Weatherstack API.
    """
    try:
        # Weatherstack API - you'll need to sign up for an API key
        api_key = "a1e511e34a3e7767793a88a0add82afb"
        base_url = "http://api.weatherstack.com/current"
        
        # Weatherstack uses 'm', 'f', or 's' for units (metric, imperial, scientific)
        unit_map = {"metric": "m", "imperial": "f"}
        unit_param = unit_map.get(unit, "m")
        
        params = {
            "access_key": api_key,
            "query": location,
            "units": unit_param
        }
        
        response = requests.get(base_url, params=params, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            if "current" in data:
                return {
                    "location": f"{data['location']['name']}, {data['location']['country']}",
                    "temperature": f"{data['current']['temperature']}°{'C' if unit == 'metric' else 'F'}",
                    "condition": data['current']['weather_descriptions'][0],
                    "humidity": f"{data['current']['humidity']}%",
                    "wind_speed": f"{data['current']['wind_speed']} {'km/h' if unit == 'metric' else 'mph'}"
                }
            else:
                return {"error": "Could not fetch weather data. Please check the location or API key."}
        else:
            return {"error": f"Could not fetch weather data. Status code: {response.status_code}"}
            
    except Exception as e:
        return {"error": str(e)}

def verify_calculations(solution_text):
    """
    Verifies critical calculations in the solution text.
    Returns corrected solution or None if no corrections needed.
    """
    try:
        # Look for calculation patterns like "X / Y = Z" or "X * Y = Z" or "X + Y = Z" or "X - Y = Z"
        calculation_patterns = [
            r'(\d+(?:\.\d+)?)\s*\/\s*(\d+(?:\.\d+)?)\s*=\s*(\d+(?:\.\d+)?)',  # Division
            r'(\d+(?:\.\d+)?)\s*\*\s*(\d+(?:\.\d+)?)\s*=\s*(\d+(?:\.\d+)?)',  # Multiplication
            r'(\d+(?:\.\d+)?)\s*\+\s*(\d+(?:\.\d+)?)\s*=\s*(\d+(?:\.\d+)?)',  # Addition
            r'(\d+(?:\.\d+)?)\s*\-\s*(\d+(?:\.\d+)?)\s*=\s*(\d+(?:\.\d+)?)'   # Subtraction
        ]
        
        corrected = solution_text
        needs_correction = False
        
        for pattern in calculation_patterns:
            for match in re.finditer(pattern, solution_text):
                a, b, claimed_result = map(float, match.groups())
                
                if '/' in match.group():
                    actual_result = a / b if b != 0 else float('inf')
                elif '*' in match.group():
                    actual_result = a * b
                elif '+' in match.group():
                    actual_result = a + b
                elif '-' in match.group():
                    actual_result = a - b
                
                # If the claimed result is significantly different from actual result
                if abs(actual_result - claimed_result) > 0.0001:
                    # Replace the incorrect calculation with correct one
                    old_calc = match.group()
                    new_calc = old_calc.replace(str(claimed_result), f"{actual_result:.4f}")
                    corrected = corrected.replace(old_calc, new_calc)
                    needs_correction = True
        
        return corrected if needs_correction else None
        
    except Exception as e:
        print(f"Verification error: {e}")
        return None
    
# Text chunking function
def chunk_text(text, chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP):
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = min(start + chunk_size, text_length)
        if end < text_length and end - start == chunk_size:
            # Find a good breaking point - like a period or newline
            for i in range(min(100, chunk_overlap)):
                if end - i > 0 and text[end - i] in ['.', '\n', '!', '?']:
                    end = end - i + 1
                    break
        
        chunks.append(text[start:end])
        
        # Move the start point, considering overlap
        if end == text_length:
            break
        start = max(start, end - chunk_overlap)
    
    return chunks

# Math expression detection
def is_math_expression(expr: str) -> bool:
    try:
        result = eval(expr)
        return isinstance(result, (int, float))
    except:
        return False

# Decide whether a question should go to SerpAPI
def is_serpapi_query(message: str) -> bool:
    sports_keywords = ["sports", "ipl", "match", "cricket", "football", "game", "tournament"]
    news_weather_keywords = ["news", "weather", "sports", "finance", "stock", "market", "update", "today", "who", "when"]
    message_lower = message.lower()
    return any(k in message_lower for k in sports_keywords + news_weather_keywords)

# Call SerpAPI with fallback key rotation
def query_serpapi(question: str) -> str:
    for key in SERPAPI_KEYS:
        params = {
            "q": question,
            "api_key": key,
            "engine": "google"
        }
        try:
            response = requests.get("https://serpapi.com/search", params=params)
            data = response.json()

            if "error" in data and ("Rate limit" in data["error"] or "You are out of searches" in data["error"]):
                print(f"API key {key} is out of quota. Trying next key...")
                continue

            if "answer_box" in data:
                if "answer" in data["answer_box"]:
                    return data["answer_box"]["answer"]
                elif "snippet" in data["answer_box"]:
                    return data["answer_box"]["snippet"]
            elif "snippet" in data.get("organic_results", [{}])[0]:
                return data["organic_results"][0]["snippet"]

            return "No relevant information found."
        except Exception as e:
            print(f"Error with API key {key}: {e}")
            continue
    return "All API keys are exhausted or failed."

# Extract text from PDF
def extract_text_from_pdf(file: UploadFile):
    reader = PyPDF2.PdfReader(file.file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Extract text from DOCX
def extract_text_from_docx(file: UploadFile):
    doc = docx.Document(file.file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Extract text from CSV
def extract_text_from_csv(file: UploadFile):
    df = pd.read_csv(file.file)
    return df.to_string()

# RAG-specific functions
def process_file_for_rag(file: UploadFile):
    # Extract text based on file type
    if file.filename.endswith(".pdf"):
        text = extract_text_from_pdf(file)
    elif file.filename.endswith(".docx"):
        text = extract_text_from_docx(file)
    elif file.filename.endswith(".csv"):
        text = extract_text_from_csv(file)
    else:
        raise ValueError("Unsupported file type")
    
    # Chunk the text
    chunks = chunk_text(text)
    
    # Create document ID
    doc_id = str(uuid4())
    
    # Create document object
    document = Document(id=doc_id, name=file.filename, chunks=chunks)
    
    # Create embeddings for each chunk
    embeddings = model.encode(chunks)
    
    # Create FAISS index
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(np.array(embeddings).astype('float32'))
    
    # Save the index and document
    index_path = os.path.join(VECTOR_DB_DIR, f"{doc_id}_index.faiss")
    document_path = os.path.join(VECTOR_DB_DIR, f"{doc_id}_document.pkl")
    
    faiss.write_index(index, index_path)
    with open(document_path, 'wb') as f:
        pickle.dump(document, f)
    
    # Update metadata
    doc_info = DocumentInfo(id=doc_id, name=file.filename)
    document_metadata.documents.append(doc_info)
    save_document_metadata(document_metadata)
    
    return doc_id

def retrieve_relevant_chunks(query: str, doc_id: str, top_k=TOP_K_RESULTS):
    # Load the index and document
    index_path = os.path.join(VECTOR_DB_DIR, f"{doc_id}_index.faiss")
    document_path = os.path.join(VECTOR_DB_DIR, f"{doc_id}_document.pkl")
    
    if not os.path.exists(index_path) or not os.path.exists(document_path):
        raise ValueError("Document not found")
    
    index = faiss.read_index(index_path)
    with open(document_path, 'rb') as f:
        document = pickle.load(f)
    
    # Create query embedding
    query_embedding = model.encode([query])[0].reshape(1, -1).astype('float32')
    
    # Search for similar chunks
    D, I = index.search(query_embedding, min(top_k, len(document.chunks)))
    
    # Get the relevant chunks
    relevant_chunks = [document.chunks[i] for i in I[0]]
    
    return relevant_chunks

# Function to build chat history for context
def build_chat_history_context(chat: Chat):
    # Get the last few messages for context (limit to prevent token overload)
    recent_messages = chat.messages[-6:] if len(chat.messages) > 6 else chat.messages
    
    context = ""
    
    # Add file context if available
    if chat.file_content:
        context += f"File: {chat.file_name}\nContent summary: {chat.file_content[:3000]}...\n\n"
    
    # Add conversation history
    context += "Previous conversation:\n"
    for msg in recent_messages:
        context += f"{msg.sender}: {msg.text}\n"
        if msg.image_path:
            context += f"[Attached visualization: {msg.image_path}]\n"
    
    return context

def save_image_to_file(base64_image, chat_id):
    """Save a base64 encoded image to a file and return the relative path"""
    try:
        # Ensure directory exists
        os.makedirs("static/images", exist_ok=True)
        
        # Generate a unique filename including chat ID for better organization
        filename = f"{chat_id}_{uuid4()}.png"
        filepath = os.path.join("static/images", filename)
        
        # Strip off data URL prefix if present
        if "," in base64_image:
            base64_data = base64_image.split(",")[1]
        else:
            base64_data = base64_image
            
        # Decode and save the image
        image_data = base64.b64decode(base64_data)
        with open(filepath, "wb") as img_file:
            img_file.write(image_data)
            
        # Return just the filename, not the full path
        return filename
    except Exception as e:
        print(f"Error saving image: {e}")
        return None

# Routes
@app.get("/languages")
def get_supported_languages():
    """Return the list of supported languages"""
    return {"languages": list(SUPPORTED_LANGUAGES.keys())}

@app.post("/new_chat")
def create_chat(document_id: Optional[str] = None, language: str = "English"):
    new_chat = Chat(id=str(uuid4()), messages=[], document_id=document_id, language=language)
    chats.append(new_chat)
    save_chats_to_file()
    return {"message": "New chat created", "chat_id": new_chat.id}

@app.delete("/chats/{chat_id}")
def delete_chat(chat_id: str):
    global chats
    chats = [chat for chat in chats if chat.id != chat_id]
    save_chats_to_file()
    return {"message": "Chat deleted successfully"}

@app.post("/upload_file")
async def upload_file(file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    try:
        # Process the file with RAG
        doc_id = process_file_for_rag(file)
        
        return {
            "document_id": doc_id,
            "filename": file.filename,
            "message": "File processed successfully for RAG"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.get("/documents")
def get_documents():
    return {"documents": document_metadata.documents}

# Modified endpoint to handle both JSON and Form data
@app.post("/chat")
async def chat_endpoint(
    message: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    document_id: Optional[str] = Form(None),
    chat_id: Optional[str] = Form(None), 
    language: str = Form("English")  # Add language parameter
):
    # Raise an error if neither a file nor a message is provided
    if message is None and file is None:
        raise HTTPException(status_code=400, detail="No input provided. Please provide a message or upload a file.")

    # Ensure we have a message
    if message is None and file:
        message = "Please analyze this file"

    # Find the specified chat or use the latest one
    current_chat = None
    if chat_id:
        current_chat = next((chat for chat in chats if chat.id == chat_id), None)
    
    # If no chat found or no chat_id provided, use the latest or create a new one
    if not current_chat:
        if not chats:
            current_chat = Chat(id=str(uuid4()), messages=[], document_id=document_id, language=language)
            chats.append(current_chat)
        else:
            current_chat = chats[-1]
    
    # Update chat language if it's different
    if current_chat.language != language:
        current_chat.language = language
    
    # Add the original message to chat history
    user_msg = Message(sender="You", text=message, language=language)
    current_chat.messages.append(user_msg)
    
    # If document_id is provided and different from the current one, create a new chat
    if document_id and current_chat.document_id != document_id:
        new_chat = Chat(id=str(uuid4()), messages=[], document_id=document_id, language=language)
        chats.append(new_chat)
        current_chat = new_chat
    
    # If the message is not in English and needs to be processed by models that need English,
    # translate it to English first
    processing_message = message
    original_language = language

    if language != "English" and language in SUPPORTED_LANGUAGES:
        try:
            # First try detecting if the message contains any non-Latin characters
            has_non_latin = any(ord(c) > 127 for c in message)
            
            # If message is already in target language script, translate directly
            if has_non_latin:
                processing_message = translate_text(message, SUPPORTED_LANGUAGES[language], "en")
            else:
                # For transliterated text, we'll let the model handle it
                # Just add language info to help the model
                processing_message = f"[This is transliterated {language}]: {message}"
            
            print(f"Processed message: {processing_message}")
        except Exception as e:
            print(f"Translation error (input): {e}")
            # Continue with original message if translation fails
            processing_message = message

    # Generate title only if it's the first user message in the chat
    if len(current_chat.messages) == 1:
        try:
            title_prompt = f"Generate strictly a 4 word title for this conversation based on this message: \"{processing_message}\""
            title_response = ollama.chat(model="mistral", messages=[{"role": "user", "content": title_prompt}])
            generated_title = title_response['message']['content'].strip().replace('"', '')
            current_chat.name = generated_title
        except Exception as e:
            current_chat.name = "Untitled Chat"

    # Process file if uploaded
    file_content = None
    if file:
        # Reset file position to the beginning
        await file.seek(0)

        if file.filename.endswith(".pdf"):
            file_content = extract_text_from_pdf(file)
        elif file.filename.endswith(".docx"):
            file_content = extract_text_from_docx(file)
        elif file.filename.endswith(".csv"):
            file_content = extract_text_from_csv(file)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type.")

        # Store file content in the chat for future context
        current_chat.file_content = file_content
        current_chat.file_name = file.filename

        # Reset file position again after reading
        await file.seek(0)

    # Get chat history context
    chat_context = build_chat_history_context(current_chat)
    
    # EXISTING CODE: Continue with the original code flow for non-tool cases
    if is_asking_name(processing_message, language):
        # Get the introduction in the appropriate language
        introduction = AI_INTRODUCTIONS.get(language, AI_INTRODUCTIONS["English"])
        
        # For languages not in our predefined introductions, translate from English
        if language not in AI_INTRODUCTIONS and language in SUPPORTED_LANGUAGES:
            introduction = translate_text(AI_INTRODUCTIONS["English"], "en", language)
            
        ai_answer = introduction
        
    elif is_math_problem(processing_message):
        print("Detected math problem - using Gemma2:9b")
        # Create math-specific prompt
        math_prompt = f"""
        You are an expert mathematics tutor. 
        Solve this math problem step-by-step, showing all work clearly:
        
        {processing_message}
        
        Follow these guidelines:
        1. Define all variables clearly
        2. Show each calculation step
        3. Double-check all arithmetic
        4. Provide the final answer with units if applicable
        5. Give Equations in LaTeX format if possible
        """
        
        # Use gemma2:9b for math problems
        response = ollama.chat(model="gemma2:9b", messages=[{"role": "user", "content": math_prompt}])
        ai_answer = response['message']['content']
        
        # Verify calculations if needed
        corrected = verify_calculations(ai_answer)
        if corrected:
            ai_answer = corrected
            
        # Translate if needed
        if language != "English" and language in SUPPORTED_LANGUAGES:
            ai_answer = translate_text(ai_answer, "en", language)
    
    # Weather queries
    elif is_weather_query(processing_message):
        location = extract_location(processing_message)
        if location:
            weather_info = get_weather_information(location)
            ai_answer = format_weather_response(weather_info, processing_message, language)
        else:
            # Fall back to regular chat if location not found
            ai_answer = await get_regular_chat_response(processing_message, chat_context, language)
            
    elif is_serpapi_query(processing_message):
        serpapi_result = query_serpapi(processing_message)
        full_prompt = f"{chat_context}\n\nUser asked: {processing_message}\nHere is some real-time information:\n{serpapi_result}\nPlease answer naturally using this context."
        try:
            response = ollama.chat(model="mistral", messages=[{"role": "user", "content": full_prompt}])
            ai_answer = response['message']['content']
            
            # Translate the answer back to the user's language if needed
            if language != "English" and language in SUPPORTED_LANGUAGES:
                ai_answer = translate_text(ai_answer, "en", language)
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI response error: {str(e)}")
            
    elif file_content:
        # If this is a new file upload
        context = file_content[:3000]  # Extended context from 1000 to 3000
        full_prompt = f"User asked: {processing_message}\nUser uploaded a file with the following content:\n{context}\nPlease answer based on this content."
        try:
            response = ollama.chat(model="mistral", messages=[{"role": "user", "content": full_prompt}])
            ai_answer = response['message']['content']
            
            # Translate the answer back to the user's language if needed
            if language != "English" and language in SUPPORTED_LANGUAGES:
                ai_answer = translate_text(ai_answer, "en", language)
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI response error: {str(e)}")
            
    elif current_chat.file_content:
        # Use existing file content for context in follow-up questions
        full_prompt = f"{chat_context}\n\nUser asked: {processing_message}\nPlease answer based on the previously uploaded file and conversation context."
        try:
            response = ollama.chat(model="mistral", messages=[{"role": "user", "content": full_prompt}])
            ai_answer = response['message']['content']
            
            # Translate the answer back to the user's language if needed
            if language != "English" and language in SUPPORTED_LANGUAGES:
                ai_answer = translate_text(ai_answer, "en", language)
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"AI response error: {str(e)}")
            
    elif current_chat.document_id:
        # RAG-based response
        try:
            # Retrieve relevant chunks
            relevant_chunks = retrieve_relevant_chunks(processing_message, current_chat.document_id)
            
            # Build context from relevant chunks
            chunks_context = "\n\n".join(relevant_chunks)
            
            # Build prompt with retrieved context and chat history
            full_prompt = f"""{chat_context}

User asked: {processing_message}
            
Here are relevant passages from the document:

{chunks_context}

Please answer the user's question based on the information provided in these passages and the conversation history. 
If the answer cannot be found in the passages, please indicate that clearly. Remember to introduce yourself as {AI_NAME} if appropriate."""
            
            # Get response from LLM
            response = ollama.chat(model="mistral", messages=[{"role": "user", "content": full_prompt}])
            ai_answer = response['message']['content']
            
            # Translate the answer back to the user's language if needed
            if language != "English" and language in SUPPORTED_LANGUAGES:
                ai_answer = translate_text(ai_answer, "en", language)
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"RAG response error: {str(e)}")
    elif any(date_word in processing_message.lower() for date_word in ["date", "today's date","current day", "what day is it"]):
        today = datetime.datetime.now()
        formatted_date = today.strftime("%A, %B %d, %Y")
        ai_answer = f"Today is {formatted_date}."
        
        # Translate the answer back to the user's language if needed
        if language != "English" and language in SUPPORTED_LANGUAGES:
            ai_answer = translate_text(ai_answer, "en", language)
    else:
        # Regular chat without file context
        try:
            # Check if this might be transliterated content
            is_transliterated = (language != "English" and all(ord(c) < 128 for c in message))
            # Include chat history for context
            full_prompt = f"{chat_context}\n\n"

            if is_transliterated:
                full_prompt += f"User asked in transliterated {language}: {message}\n"
                full_prompt += f"Please understand this {language} request written in Latin script and respond in formal {language}.\n"
            else:
                full_prompt += f"User asked: {processing_message}\n"
            

            # Include chat history for context
            print("Using Mistral for regular conversation")
            full_prompt = f"{chat_context}\n\nUser asked: {processing_message}\nPlease respond based on the conversation context. Remember to introduce yourself as {AI_NAME} if appropriate."
            response = ollama.chat(model="mistral", messages=[{"role": "user", "content": full_prompt}])
            ai_answer = response['message']['content']

            # Wrap response in code block if it's code
            if any(keyword in processing_message.lower() for keyword in ["code", "function", "script", "program", "class", "write", "implement"]):
                if "```" not in ai_answer:
                    ai_answer = f"```python\n{ai_answer.strip()}\n```"
            
            # Translate the answer back to the user's language if needed
            if language != "English" and language in SUPPORTED_LANGUAGES:
                # For transliterated input, check if the model already responded in the target language script
                has_target_script = False
                if is_transliterated:
                    # Check if response already contains non-Latin characters that might be the target language
                    has_target_script = any(ord(c) > 127 for c in ai_answer)
                
                if not has_target_script:
                    try:
                        if "```" in ai_answer:
                            # Split by code blocks and translate only the non-code parts
                            parts = ai_answer.split("```")
                            for i in range(0, len(parts), 2):  # Only translate even indices (non-code parts)
                                parts[i] = translate_text(parts[i], "en", SUPPORTED_LANGUAGES[language])
                            ai_answer = "```".join(parts)
                        else:
                            ai_answer = translate_text(ai_answer, "en", SUPPORTED_LANGUAGES[language])
                    except Exception as e:
                        print(f"Translation error (output): {e}")
                        # If translation fails, ensure we at least inform the user
                        if all(ord(c) < 128 for c in ai_answer):  # If answer is in Latin script
                            ai_answer = f"{ai_answer}\n\n(Translation to {language} failed, showing English response)"
        except Exception as e:
            error_msg = f"AI response error: {str(e)}"
            print(error_msg)
            ai_answer = f"I apologize, but I encountered an error while processing your request. Please try again."
            if language != "English" and language in SUPPORTED_LANGUAGES:
                try:
                    ai_answer = translate_text(ai_answer, "en", SUPPORTED_LANGUAGES[language])
                except:
                    pass 

    ai_msg = Message(sender="AI", text=ai_answer, language=language)
    current_chat.messages.append(ai_msg)
    save_chats_to_file()
    print(ai_answer)
    return {"response": ai_answer, "chat_id": current_chat.id}

@app.get("/chats")
def get_chats():
    return {"chats": chats}

# New endpoints for RAG functionality
@app.post("/chat_with_document/{document_id}")
async def chat_with_document(document_id: str, request: ChatRequest):
    try:
        # Process message in the requested language
        processing_message = request.message
        if request.language != "English" and request.language in SUPPORTED_LANGUAGES:
            # Translate to English for processing
            processing_message = translate_text(request.message, request.language, "en")
        
        # Check if the user is asking for the AI's name
        if is_asking_name(processing_message, request.language):
            # Get the introduction in the appropriate language
            introduction = AI_INTRODUCTIONS.get(request.language, AI_INTRODUCTIONS["English"])
            
            # For languages not in our predefined introductions, translate from English
            if request.language not in AI_INTRODUCTIONS and request.language in SUPPORTED_LANGUAGES:
                introduction = translate_text(AI_INTRODUCTIONS["English"], "en", request.language)
                
            # Find or create a chat for this document
            document_chat = next((chat for chat in chats if chat.document_id == document_id), None)
            
            if document_chat is None:
                document_chat = Chat(id=str(uuid4()), messages=[], document_id=document_id, language=request.language)
                chats.append(document_chat)
            
            # Add messages to chat
            user_msg = Message(sender="You", text=request.message, language=request.language)
            ai_msg = Message(sender="AI", text=introduction, language=request.language)
            document_chat.messages.extend([user_msg, ai_msg])
            
            # Save chat history
            save_chats_to_file()
            
            return {"response": introduction, "chat_id": document_chat.id}
        
        # Retrieve relevant chunks
        relevant_chunks = retrieve_relevant_chunks(processing_message, document_id)
        
        # Build context from relevant chunks
        context = "\n\n".join(relevant_chunks)
        
        # Find or create a chat for this document
        document_chat = next((chat for chat in chats if chat.document_id == document_id), None)
        
        if document_chat is None:
            document_chat = Chat(id=str(uuid4()), messages=[], document_id=document_id, language=request.language)
            chats.append(document_chat)
        
        # Get chat history context
        chat_context = build_chat_history_context(document_chat)
        
        # Build prompt with retrieved context and chat history
        full_prompt = f"""{chat_context}

User asked: {processing_message}
        
Here are relevant passages from the document:

{context}

Please answer the user's question based on the information provided in these passages and the conversation history. 
If the answer cannot be found in the passages, please indicate that clearly. Remember to introduce yourself as {AI_NAME} if appropriate."""
        
        # Get response from LLM
        response = ollama.chat(model="mistral", messages=[{"role": "user", "content": full_prompt}])
        ai_answer = response['message']['content']
        
        # Translate the answer back to the user's language if needed
        if request.language != "English" and request.language in SUPPORTED_LANGUAGES:
            ai_answer = translate_text(ai_answer, "en", request.language)
        
        # Add messages to chat
        user_msg = Message(sender="You", text=request.message, language=request.language)
        ai_msg = Message(sender="AI", text=ai_answer, language=request.language)
        document_chat.messages.extend([user_msg, ai_msg])
        
        # Save chat history
        save_chats_to_file()
        
        return {"response": ai_answer, "chat_id": document_chat.id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")
    
@app.post("/visualize")
async def visualize_data(request: VisualizationRequest):
    try:
        # Get the chat if a chat_id is provided
        chat = None
        if request.chat_id:
            chat = next((c for c in chats if c.id == request.chat_id), None)
            if not chat:
                raise HTTPException(status_code=404, detail="Chat not found")
        
        # Process message in the requested language
        processing_message = request.query
        if request.language != "English" and request.language in SUPPORTED_LANGUAGES:
            # Translate to English for processing
            processing_message = translate_text(request.query, request.language, "en")
        
        # Build full prompt for data extraction
        if chat:
            chat_context = build_chat_history_context(chat)
            prompt = f"""{chat_context}

Based on our conversation, I need data for visualization. 
Query: {processing_message}

Please provide data in a structured format that can be used for a {request.chart_type} chart.
Format your response as a table or list with categorical labels and numerical values.
"""
        else:
            # If no chat_id was provided, create a new chat
            chat = Chat(id=str(uuid4()), messages=[], language=request.language)
            chats.append(chat)
            prompt = f"""I need data for visualization.
Query: {processing_message}

Please provide data in a structured format that can be used for a {request.chart_type} chart.
Format your response as a table or list with categorical labels and numerical values.
"""

        # If document_id is provided, add relevant document chunks to the context
        if request.document_id:
            try:
                relevant_chunks = retrieve_relevant_chunks(processing_message, request.document_id)
                chunks_context = "\n\n".join(relevant_chunks)
                prompt += f"\n\nHere are relevant passages from the document:\n\n{chunks_context}"
                
                # Update the chat with the document_id if it's a new chat
                if not chat.document_id:
                    chat.document_id = request.document_id
            except Exception as e:
                print(f"Error retrieving document chunks: {e}")
        
        # Get data from AI model
        try:
            # Use mistral for data generation - it handles structured data well
            response = ollama.chat(model="mistral", messages=[{"role": "user", "content": prompt}])
            ai_response = response['message']['content']
            
            # Extract data for visualization from the AI response
            extracted_data = extract_data_for_visualization(ai_response)
            
            if not extracted_data or len(extracted_data["labels"]) < 2:
                # If extraction failed or not enough data points, ask specifically for tabular data
                refined_prompt = f"{prompt}\n\nPlease provide the data in a clear tabular format with labels and values, like this:\nLabel1: 10\nLabel2: 20\nLabel3: 30"
                response = ollama.chat(model="mistral", messages=[{"role": "user", "content": refined_prompt}])
                ai_response = response['message']['content']
                extracted_data = extract_data_for_visualization(ai_response)
            
            if not extracted_data or len(extracted_data["labels"]) < 2:
                return {"error": "Could not extract visualization data from the AI response"}
            
            # Generate visualization
            image_data = generate_visualization(extracted_data, request.chart_type)
            if not image_data:
                return {"error": "Failed to generate visualization"}
            
            # Create visualization data response
            viz_data = VisualizationData(
                title=f"{request.chart_type.capitalize()} Chart",
                data=[DataPoint(label=label, value=value) for label, value in zip(extracted_data["labels"], extracted_data["values"])],
                chart_type=request.chart_type,
                x_label="Categories",
                y_label="Values"
            )

            image_path = save_image_to_file(image_data, chat.id)
            
            # Create visualization response with both raw data and image
            response_data = {
            "visualization_data": viz_data.dict(),
            "image": f"data:image/png;base64,{image_data}",
            "image_path": image_path,  # Include path to saved image
            "raw_response": ai_response,
            "chat_id": chat.id
            }
            
            # Add the user's query to the chat history first
            user_msg = Message(sender="You", text=request.query, language=request.language)
            chat.messages.append(user_msg)
            
            # Add visualization information and data to the AI response
            viz_message = f"Here's a {request.chart_type} chart visualization based on your request."
            if request.language != "English" and request.language in SUPPORTED_LANGUAGES:
                viz_message = translate_text(viz_message, "en", request.language)
            
            # Add AI message about the visualization
            ai_msg = Message(
            sender="AI", 
            text=viz_message, 
            language=request.language,
            image_path=image_path  # Store reference to the image
            )
            chat.messages.append(ai_msg)
            
            # Generate a title for the chat if it's new
            if chat.name == "Untitled Chat" or not chat.name:
                try:
                    title_prompt = f"Generate strictly a 4 word title for this conversation based on this message: \"{request.query}\""
                    title_response = ollama.chat(model="mistral", messages=[{"role": "user", "content": title_prompt}])
                    generated_title = title_response['message']['content'].strip().replace('"', '')
                    chat.name = generated_title
                except Exception as e:
                    chat.name = "Visualization Chat"
            
            # Save chat history
            save_chats_to_file()
            
            return response_data
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error getting data from AI model: {str(e)}")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Visualization error: {str(e)}")
    
@app.post("/generate_image")
def generate_image(prompt: str = Body(..., embed=True)):
    try:
        # Create or find a chat for this prompt
        chat = Chat(id=str(uuid4()), messages=[], language="English")
        chats.append(chat)

        # Generate image using Stable Diffusion
        with torch.autocast("cuda" if torch.cuda.is_available() else "cpu"):
            result = image_pipe(
                prompt,
                guidance_scale=8.5,  # Slightly stronger guidance for clearer outputs
                num_inference_steps=50,  # More steps for better quality
                negative_prompt="blur, low quality, distorted, duplicate, unnatural"
            )
        generated_image = result.images[0]

        # Save image to file
        buffer = io.BytesIO()
        generated_image.save(buffer, format="PNG")
        base64_image = base64.b64encode(buffer.getvalue()).decode("utf-8")

        image_path = save_image_to_file(base64_image, chat.id)

        # Save prompt and image into chat
        user_msg = Message(sender="You", text=prompt, language="English")
        chat.messages.append(user_msg)

        ai_msg = Message(
            sender="AI",
            text="Here's an image generated based on your prompt.",
            language="English",
            image_path=image_path
        )
        chat.messages.append(ai_msg)

        # Generate chat title
        try:
            title_prompt = f"Generate strictly a 4 word title for this conversation based on this prompt: \"{prompt}\""
            response = ollama.chat(model="mistral", messages=[{"role": "user", "content": title_prompt}])
            chat.name = response['message']['content'].strip().replace('"', '')
        except:
            chat.name = "Generated Image"

        save_chats_to_file()

        return {
            "image_path": image_path,
            "chat_id": chat.id,
            "chat_name": chat.name
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    
@app.get("/images/{filename}")
async def get_image(filename: str):
    file_path = f"static/images/{filename}"
    if not os.path.exists(file_path):
        return Response(status_code=404)
    headers = {
        "Access-Control-Allow-Origin": "*"
    }
    return FileResponse(file_path, headers=headers)

@app.get("/chats/{chat_id}")
async def get_chat(chat_id: str):
    # Make sure to include any visualization data or image paths in the messages
    # This ensures all necessary data is available when loading a chat
    chat = next((c for c in chats if c.id == chat_id), None)
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    chat_data = {
        "id": chat.id,
        "name": chat.name,
        "messages": [
            {
                "sender": msg.sender,
                "text": msg.text,
                "language": msg.language,
                "image_path": msg.image_path if hasattr(msg, 'image_path') else None
            }
            for msg in chat.messages
        ],
        "document_id": chat.document_id,
        "file_content": chat.file_content,
        "file_name": chat.file_name
    }
    return chat_data